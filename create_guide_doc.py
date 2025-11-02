from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import datetime

# 문서 생성
doc = Document()

# 문서 여백 설정
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# 제목 추가
title = doc.add_heading('대규모 프로젝트 개발 가이드', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 부제목
subtitle = doc.add_paragraph()
subtitle.add_run('비개발자를 위한 체계적인 LMS 개발 관리 방법').bold = True
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph()

# 작성일
date_para = doc.add_paragraph()
date_para.add_run(f'작성일: {datetime.datetime.now().strftime("%Y년 %m월 %d일")}')
date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
doc.add_paragraph()

# 1. 문제점 정리
doc.add_heading('1. 현재 겪고 계신 문제점', 1)
doc.add_paragraph('대규모 프로젝트를 진행하시면서 이런 어려움을 겪고 계시죠?')
doc.add_paragraph()

problems = [
    '새로운 기능을 추가하면 기존에 잘 되던 기능이 갑자기 안 됨',
    'Claude와 대화가 길어지면 앞에서 만든 코드를 잊어버림',
    '새 대화창에서 시작하면 프로젝트를 처음부터 다시 설명해야 함',
    '어디까지 개발했는지 헷갈리고 정리가 안 됨'
]

for problem in problems:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(problem)

doc.add_paragraph()
doc.add_paragraph('이 문제들을 해결하기 위한 실용적인 방법을 알려드리겠습니다.')

# 2. 핵심 해결책
doc.add_heading('2. 핵심 해결책: 레고 블록처럼 나누어 개발하기', 1)
doc.add_paragraph()
doc.add_paragraph('큰 건물을 한 번에 짓는 것보다 레고 블록을 하나씩 조립하는 것이 쉬운 것처럼, ')
doc.add_paragraph('LMS도 작은 부분으로 나누어 개발하면 훨씬 관리하기 쉬워집니다.')
doc.add_paragraph()

# 폴더 구조 예시
doc.add_heading('2-1. 폴더 정리하기 (서랍장 정리하듯이)', 2)
doc.add_paragraph()
doc.add_paragraph('프로젝트 폴더를 이렇게 정리해보세요:')
doc.add_paragraph()

# 표 추가
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '폴더 이름'
hdr_cells[1].text = '무엇을 넣을까요?'

folder_structure = [
    ('📁 핵심기능', '로그인, 회원가입 등 모든 곳에서 쓰는 기능'),
    ('📁 학생관리', '학생 등록, 수정, 삭제 관련 기능'),
    ('📁 수업관리', '수업 생성, 시간표 관련 기능'),
    ('📁 출석관리', '출석 체크, 출석부 관련 기능'),
    ('📁 성적관리', '시험 점수, 성적표 관련 기능'),
    ('📁 백업폴더', '중요한 시점의 코드 백업본 보관')
]

for folder, description in folder_structure:
    row_cells = table.add_row().cells
    row_cells[0].text = folder
    row_cells[1].text = description

doc.add_paragraph()

# 3. 실전 가이드
doc.add_heading('3. 따라하기 쉬운 실전 가이드', 1)
doc.add_paragraph()

# Step 1
doc.add_heading('STEP 1: 프로젝트 상태 기록장 만들기', 2)
doc.add_paragraph()
doc.add_paragraph('메모장을 열고 "프로젝트_현황.txt" 파일을 만들어 이렇게 작성하세요:')
doc.add_paragraph()

# 코드 블록 스타일로 예시
example = doc.add_paragraph()
example.add_run('==== 우리 LMS 프로젝트 현황 ====\n').bold = True
example.add_run('\n')
example.add_run('🟢 완성된 기능:\n')
example.add_run('  - 로그인/로그아웃 ✓\n')
example.add_run('  - 학생 등록 ✓\n')
example.add_run('\n')
example.add_run('🟡 개발 중인 기능:\n')
example.add_run('  - 출석 체크 기능\n')
example.add_run('\n')
example.add_run('🔴 아직 안 만든 기능:\n')
example.add_run('  - 성적 관리\n')
example.add_run('  - 학부모 알림\n')
example.add_run('\n')
example.add_run('⚠️ 절대 수정하면 안 되는 것:\n')
example.add_run('  - login.js 파일\n')
example.add_run('  - database 설정 파일\n')
example.style = 'Quote'

doc.add_paragraph()

# Step 2
doc.add_heading('STEP 2: Claude와 효율적으로 대화하기', 2)
doc.add_paragraph()
doc.add_paragraph('새로운 기능을 만들 때마다 이렇게 시작하세요:')
doc.add_paragraph()

# 대화 예시 박스
conv_example = doc.add_paragraph()
conv_example.add_run('대화 시작 예시:\n').bold = True
conv_example.add_run('─────────────────────────\n')
conv_example.add_run('"안녕 Claude, 학원 LMS 프로젝트를 진행 중이야.\n')
conv_example.add_run('현재 상황:\n')
conv_example.add_run('- 로그인 기능 완성됨\n')
conv_example.add_run('- 학생 관리 완성됨\n')
conv_example.add_run('- 지금 출석 관리 기능만 추가하고 싶어\n')
conv_example.add_run('- 다른 기능은 절대 건드리지 마\n')
conv_example.add_run('여기 출석 관리 폴더의 코드야: [코드 붙여넣기]"')
conv_example.style = 'Quote'

doc.add_paragraph()

# Step 3
doc.add_heading('STEP 3: 체크리스트로 검증하기', 2)
doc.add_paragraph()
doc.add_paragraph('새 기능을 추가한 후에는 반드시 이것들을 확인하세요:')
doc.add_paragraph()

checklist = [
    '로그인이 여전히 잘 되나요?',
    '학생 목록이 제대로 보이나요?',
    '이전에 만든 메뉴들이 모두 작동하나요?',
    '화면이 깨지지 않았나요?',
    '에러 메시지가 뜨지 않나요?'
]

for item in checklist:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('□ ' + item)

doc.add_paragraph()

# 4. 백업 전략
doc.add_heading('4. 안전한 백업 전략', 1)
doc.add_paragraph()
doc.add_paragraph('중요한 작업 전에는 반드시 백업하세요!')
doc.add_paragraph()

# 백업 폴더 구조
doc.add_heading('백업 폴더 만들기', 2)
backup_para = doc.add_paragraph()
backup_para.add_run('백업 폴더 예시:\n').bold = True
backup_para.add_run('📁 백업_2024_11_20_출석전\n')
backup_para.add_run('📁 백업_2024_11_21_성적전\n')
backup_para.add_run('📁 백업_2024_11_22_최종완성\n')
backup_para.style = 'Quote'

doc.add_paragraph()
doc.add_paragraph('이렇게 하면 문제가 생겼을 때 언제든 이전 버전으로 돌아갈 수 있습니다.')

# 5. 작업 순서
doc.add_heading('5. 추천하는 작업 순서', 1)
doc.add_paragraph()

work_order = [
    ('계획', '오늘 만들 기능 1개만 정하기'),
    ('백업', '현재 코드를 날짜별 폴더에 복사'),
    ('Claude 대화 시작', '프로젝트 현황과 오늘 할 일만 설명'),
    ('개발', '한 번에 하나씩만 추가'),
    ('테스트', '체크리스트로 기존 기능 확인'),
    ('기록', '프로젝트_현황.txt 업데이트')
]

for i, (title, desc) in enumerate(work_order, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}단계. {title}: ').bold = True
    p.add_run(desc)

doc.add_paragraph()

# 6. 문제 해결 팁
doc.add_heading('6. 자주 발생하는 문제와 해결법', 1)
doc.add_paragraph()

# 문제-해결 표
problem_table = doc.add_table(rows=1, cols=2)
problem_table.style = 'Light List Accent 1'
hdr = problem_table.rows[0].cells
hdr[0].text = '문제 상황'
hdr[1].text = '해결 방법'

problems_solutions = [
    ('Claude가 이전 내용을 잊어버려요', 
     '프로젝트_현황.txt 내용을 복사해서 대화 시작할 때마다 보여주세요'),
    ('새 기능 추가했더니 기존 기능이 안 돼요', 
     '백업 폴더에서 이전 버전을 복원하고 다시 시도하세요'),
    ('어디까지 했는지 모르겠어요', 
     '프로젝트_현황.txt를 매일 업데이트하세요'),
    ('Claude가 전체 코드를 다 수정해버려요', 
     '"출석 폴더의 파일만 수정해줘" 라고 명확히 요청하세요'),
    ('대화가 너무 길어졌어요', 
     '현재 작업을 마무리하고 새 대화창에서 다음 기능 시작하세요')
]

for problem, solution in problems_solutions:
    row = problem_table.add_row().cells
    row[0].text = problem
    row[1].text = solution

doc.add_paragraph()

# 7. Claude 활용 템플릿
doc.add_heading('7. Claude에게 복사-붙여넣기할 템플릿', 1)
doc.add_paragraph()
doc.add_paragraph('아래 템플릿을 복사해서 사용하세요:')
doc.add_paragraph()

template = doc.add_paragraph()
template.add_run('=== 새 기능 개발 요청 템플릿 ===\n\n').bold = True
template.add_run('프로젝트: 학원 LMS 시스템\n')
template.add_run('기술: React, Node.js, MySQL\n\n')
template.add_run('✅ 완성된 기능:\n')
template.add_run('- [완성된 기능 리스트]\n\n')
template.add_run('🎯 오늘 만들 기능:\n')
template.add_run('- [한 가지 기능만 작성]\n\n')
template.add_run('📁 작업할 폴더:\n')
template.add_run('- [폴더명] 폴더만 수정\n\n')
template.add_run('⚠️ 주의사항:\n')
template.add_run('- 다른 폴더의 파일은 수정하지 마세요\n')
template.add_run('- 기존 로그인 기능은 그대로 유지해주세요\n\n')
template.add_run('현재 코드:\n')
template.add_run('[해당 폴더의 코드 붙여넣기]')
template.style = 'Quote'

doc.add_paragraph()

# 8. 성공 전략 요약
doc.add_heading('8. 성공적인 개발을 위한 황금 규칙', 1)
doc.add_paragraph()

golden_rules = [
    '한 번에 한 가지 기능만 만들기',
    '매일 백업하기',
    '프로젝트 현황 문서 업데이트하기',
    'Claude에게 명확한 범위 지정하기',
    '체크리스트로 검증하기',
    '문제 생기면 백업에서 복원하기'
]

for i, rule in enumerate(golden_rules, 1):
    p = doc.add_paragraph()
    p.add_run(f'규칙 {i}. ').bold = True
    p.add_run(rule)
    
doc.add_paragraph()

# 마무리
doc.add_heading('마치며', 1)
doc.add_paragraph()
doc.add_paragraph('이 가이드를 따라 하시면 대규모 프로젝트도 체계적으로 관리할 수 있습니다.')
doc.add_paragraph('처음에는 번거로워 보일 수 있지만, 이렇게 하면 오히려 시간을 절약하고 ')
doc.add_paragraph('스트레스를 줄일 수 있습니다.')
doc.add_paragraph()
doc.add_paragraph('프로젝트 진행하시면서 궁금한 점이 있으면 언제든 물어보세요!')
doc.add_paragraph()
doc.add_paragraph('화이팅! 🎯')

# 문서 저장
import os
output_path = os.path.join(os.getcwd(), 'LMS_개발_가이드.docx')
doc.save(output_path)

print(f"문서가 성공적으로 생성되었습니다: {output_path}")
